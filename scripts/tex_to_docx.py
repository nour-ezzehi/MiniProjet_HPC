#!/usr/bin/env python3
"""
Simple LaTeX -> DOCX converter tailored for this project's `rapport.tex`.
It handles sections, subsections, itemize/enumerate, verbatim/lstlisting blocks and plain paragraphs.
Not a full LaTeX parser but sufficient to produce a readable Word document from the provided file.
"""
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    print("python-docx not installed. Please run: pip3 install python-docx")
    raise

INPUT = Path(__file__).parents[1] / 'rapport.tex'
OUTPUT = Path(__file__).parents[1] / 'rapport.docx'

text = INPUT.read_text(encoding='utf-8')

# Extract content between \begin{document} and \end{document}
m = re.search(r"\\begin\{document\}(.*)\\end\{document\}", text, re.S)
content = m.group(1) if m else text

# Remove comments
content = re.sub(r"(?m)^%.*$", "", content)

# Normalize some spacing
content = content.replace('\\\n', '\n')

# Helpers
def strip_tex_commands(s):
    # remove inline commands like \emph{...}, \textbf{...}, \texttt{...}
    s = re.sub(r"\\(emph|textbf|texttt)\{([^}]*)\}", r"\2", s)
    # remove simple commands like \item[] or \vspace{...}
    s = re.sub(r"\\[a-zA-Z]+\*?(\[[^\]]*\])?(\{[^}]*\})?", "", s)
    # remove math $...$
    s = re.sub(r"\$[^$]*\$", "", s)
    # remove remaining curly braces
    s = s.replace('{', '').replace('}', '')
    return s

# Split content into lines for processing
lines = content.splitlines()

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

i = 0
n = len(lines)

while i < n:
    line = lines[i].strip()
    if not line:
        i += 1
        continue

    # Section
    m = re.match(r"\\section\{(.+)\}", line)
    if m:
        doc.add_heading(strip_tex_commands(m.group(1)), level=1)
        i += 1
        continue
    m = re.match(r"\\subsection\{(.+)\}", line)
    if m:
        doc.add_heading(strip_tex_commands(m.group(1)), level=2)
        i += 1
        continue
    m = re.match(r"\\subsubsection\{(.+)\}", line)
    if m:
        doc.add_heading(strip_tex_commands(m.group(1)), level=3)
        i += 1
        continue

    # begin{itemize} or enumerate
    if line.startswith('\\begin{itemize}') or line.startswith('\\begin{enumerate}'):
        list_type = 'bullet' if 'itemize' in line else 'number'
        i += 1
        while i < n and not lines[i].strip().startswith('\\end{itemize}') and not lines[i].strip().startswith('\\end{enumerate}'):
            l = lines[i].strip()
            if l.startswith('\\item'):
                item_text = l[5:].strip()
                # if the item continues on following lines, gather them
                j = i + 1
                cont = []
                while j < n and not lines[j].strip().startswith('\\item') and not lines[j].strip().startswith('\\end{itemize}') and not lines[j].strip().startswith('\\end{enumerate}'):
                    cont.append(lines[j].strip())
                    j += 1
                if cont:
                    item_text += ' ' + ' '.join(cont)
                p = doc.add_paragraph(strip_tex_commands(item_text))
                if list_type == 'bullet':
                    p.style = 'List Bullet'
                else:
                    p.style = 'List Number'
                i = j
            else:
                i += 1
        i += 1
        continue

    # verbatim or lstlisting
    if line.startswith('\\begin{verbatim}') or line.startswith('\\begin{lstlisting}'):
        block_lines = []
        i += 1
        while i < n and not (lines[i].strip().startswith('\\end{verbatim}') or lines[i].strip().startswith('\\end{lstlisting}')):
            block_lines.append(lines[i].rstrip())
            i += 1
        # Add as a preformatted paragraph
        for bl in block_lines:
            p = doc.add_paragraph(bl)
            p.style.font.name = 'Courier New'
        i += 1
        continue

    # plain paragraph: gather until empty line or LaTeX command line
    para_lines = [line]
    j = i + 1
    while j < n and lines[j].strip() and not re.match(r"\\(section|subsection|subsubsection|begin)\b", lines[j].strip()):
        para_lines.append(lines[j].strip())
        j += 1
    paragraph = ' '.join(para_lines)
    paragraph = strip_tex_commands(paragraph)
    if paragraph:
        doc.add_paragraph(paragraph)
    i = j

# Save
doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")

